"""
app/service/export/router.py — 소설 내보내기 엔드포인트

POST /export/download
  - format: "txt" | "md" | "docx" | "pdf" | "epub"
  - novel_title, author, episodes: [{episode_no, title, content_html}]
  - 파일을 StreamingResponse로 반환
"""
from __future__ import annotations

import io
import os
import re
import textwrap
from datetime import date
from typing import List, Optional

from bs4 import BeautifulSoup
from fastapi import APIRouter
from fastapi.responses import StreamingResponse
from pydantic import BaseModel

router = APIRouter(prefix="/export", tags=["Export"])

# ── 한국어 폰트 경로 (PDF 전용) ────────────────────────────────────────────────
_FONT_CANDIDATES = [
    r"C:\Windows\Fonts\malgun.ttf",       # Windows – 맑은 고딕
    r"C:\Windows\Fonts\HANBatang.ttf",    # Windows – HAN 바탕
    "/usr/share/fonts/truetype/nanum/NanumGothic.ttf",    # Ubuntu/Debian
    "/usr/share/fonts/nanum/NanumGothic.ttf",
    "/System/Library/Fonts/AppleSDGothicNeo.ttc",         # macOS
]


def _find_korean_font() -> Optional[str]:
    for p in _FONT_CANDIDATES:
        if os.path.exists(p):
            return p
    return None


# ── 요청 스키마 ────────────────────────────────────────────────────────────────

class EpisodeItem(BaseModel):
    episode_no: int
    title: str
    content_html: str


class ExportRequest(BaseModel):
    format: str                    # "txt" | "md" | "docx" | "pdf" | "epub"
    novel_title: str = "소설"
    author: str = ""
    episodes: List[EpisodeItem] = []


# ── HTML → 순수 텍스트 ────────────────────────────────────────────────────────

def _html_to_text(html: str, keep_paragraphs: bool = True) -> str:
    """
    TipTap HTML → 순수 텍스트
    - <p> → 줄바꿈 두 번
    - <br> → 줄바꿈 한 번
    - <strong>/<em> 등 인라인 태그 제거
    """
    soup = BeautifulSoup(html or "", "html.parser")

    # <br> → 개행 마커
    for br in soup.find_all("br"):
        br.replace_with("\n")

    # <p>, <h1>~<h6>, <blockquote>, <li> → 단락 구분
    block_tags = ["p", "h1", "h2", "h3", "h4", "h5", "h6", "blockquote", "li", "div"]
    for tag in soup.find_all(block_tags):
        tag.insert_after("\n\n" if keep_paragraphs else "\n")

    text = soup.get_text()
    # 연속 공백/개행 정리
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


# ── HTML → Markdown ───────────────────────────────────────────────────────────

def _html_to_markdown(html: str) -> str:
    soup = BeautifulSoup(html or "", "html.parser")
    parts: list[str] = []

    def _walk(node) -> str:
        if isinstance(node, str):
            return node

        tag = node.name if hasattr(node, "name") else None
        children = "".join(_walk(c) for c in node.children)

        if tag in ("h1",):          return f"# {children}\n\n"
        if tag in ("h2",):          return f"## {children}\n\n"
        if tag in ("h3",):          return f"### {children}\n\n"
        if tag in ("strong", "b"):  return f"**{children}**"
        if tag in ("em", "i"):      return f"*{children}*"
        if tag in ("u",):           return f"<u>{children}</u>"
        if tag in ("s", "del"):     return f"~~{children}~~"
        if tag == "br":             return "\n"
        if tag == "p":              return f"{children}\n\n"
        if tag == "blockquote":     return f"> {children.strip()}\n\n"
        if tag == "li":             return f"- {children.strip()}\n"
        if tag in ("ul", "ol"):     return f"{children}\n"
        if tag == "hr":             return "---\n\n"
        return children

    for child in soup.children:
        parts.append(_walk(child))

    md = "".join(parts)
    md = re.sub(r"\n{3,}", "\n\n", md)
    return md.strip()


# ── TXT 생성 ──────────────────────────────────────────────────────────────────

def _build_txt(req: ExportRequest) -> bytes:
    lines: list[str] = []
    lines.append(req.novel_title)
    if req.author:
        lines.append(f"저자: {req.author}")
    lines.append("")
    lines.append("=" * 40)
    lines.append("")

    for ep in req.episodes:
        lines.append(f"제{ep.episode_no}화  {ep.title}")
        lines.append("-" * 30)
        lines.append(_html_to_text(ep.content_html))
        lines.append("")
        lines.append("")

    return "\n".join(lines).encode("utf-8")


# ── Markdown 생성 ─────────────────────────────────────────────────────────────

def _build_md(req: ExportRequest) -> bytes:
    lines: list[str] = []
    lines.append(f"# {req.novel_title}")
    if req.author:
        lines.append(f"**저자**: {req.author}")
    lines.append("")
    lines.append("---")
    lines.append("")

    for ep in req.episodes:
        lines.append(f"## 제{ep.episode_no}화  {ep.title}")
        lines.append("")
        lines.append(_html_to_markdown(ep.content_html))
        lines.append("")
        lines.append("---")
        lines.append("")

    return "\n".join(lines).encode("utf-8")


# ── DOCX 생성 ─────────────────────────────────────────────────────────────────

def _build_docx(req: ExportRequest) -> bytes:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # 기본 스타일 – 맑은 고딕
    style = doc.styles["Normal"]
    style.font.name = "맑은 고딕"
    style.font.size = Pt(11)

    # 표지 페이지
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run(req.novel_title)
    run.bold = True
    run.font.size = Pt(22)
    run.font.name = "맑은 고딕"

    if req.author:
        author_para = doc.add_paragraph()
        author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        ar = author_para.add_run(req.author)
        ar.font.size = Pt(13)
        ar.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
        ar.font.name = "맑은 고딕"

    doc.add_paragraph()
    doc.add_page_break()

    # 각 회차
    for ep in req.episodes:
        # 회차 제목
        h = doc.add_heading(f"제{ep.episode_no}화  {ep.title}", level=2)
        h.runs[0].font.name = "맑은 고딕"
        doc.add_paragraph()

        # 본문 파싱
        soup = BeautifulSoup(ep.content_html or "", "html.parser")
        for tag in soup.find_all(["p", "h1", "h2", "h3", "blockquote", "li"]):
            para = doc.add_paragraph()

            if tag.name in ("h1", "h2", "h3"):
                para.style = doc.styles["Heading 3"]
            elif tag.name == "blockquote":
                para.paragraph_format.left_indent = Pt(24)

            for child in tag.children:
                if isinstance(child, str):
                    run = para.add_run(child)
                    run.font.name = "맑은 고딕"
                elif hasattr(child, "name"):
                    text = child.get_text()
                    run = para.add_run(text)
                    run.font.name = "맑은 고딕"
                    if child.name in ("strong", "b"):
                        run.bold = True
                    if child.name in ("em", "i"):
                        run.italic = True
                    if child.name == "u":
                        run.underline = True

        doc.add_paragraph()

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── PDF 생성 ──────────────────────────────────────────────────────────────────

def _build_pdf(req: ExportRequest) -> bytes:
    from fpdf import FPDF

    font_path = _find_korean_font()

    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=20)
    pdf.set_margins(left=25, top=25, right=25)

    # 한국어 폰트 등록
    if font_path:
        pdf.add_font("Korean", "", font_path, uni=True)
        pdf.add_font("Korean", "B", font_path, uni=True)  # Bold fallback to same
        font_name = "Korean"
    else:
        font_name = "Helvetica"  # Korean may not render

    # ── 표지 ──
    pdf.add_page()
    pdf.set_font(font_name, size=24)
    pdf.ln(30)
    pdf.multi_cell(0, 12, req.novel_title, align="C")
    pdf.ln(8)

    if req.author:
        pdf.set_font(font_name, size=13)
        pdf.set_text_color(90, 90, 90)
        pdf.multi_cell(0, 8, req.author, align="C")
        pdf.set_text_color(0, 0, 0)

    pdf.ln(10)
    pdf.set_font(font_name, size=10)
    pdf.set_text_color(160, 160, 160)
    pdf.multi_cell(0, 6, str(date.today().year), align="C")
    pdf.set_text_color(0, 0, 0)

    # ── 본문 ──
    for ep in req.episodes:
        pdf.add_page()

        # 회차 제목
        pdf.set_font(font_name, size=16)
        pdf.multi_cell(0, 10, f"제{ep.episode_no}화  {ep.title}", align="L")
        pdf.set_draw_color(180, 180, 180)
        pdf.line(pdf.l_margin, pdf.get_y() + 2, pdf.w - pdf.r_margin, pdf.get_y() + 2)
        pdf.ln(8)

        # 본문 텍스트
        text = _html_to_text(ep.content_html)
        pdf.set_font(font_name, size=11)

        for paragraph in text.split("\n\n"):
            paragraph = paragraph.strip()
            if not paragraph:
                continue
            # 들여쓰기(첫 줄 2em 공백)
            pdf.multi_cell(0, 7, "  " + paragraph, align="J")
            pdf.ln(3)

    buf = io.BytesIO()
    pdf.output(buf)
    return buf.getvalue()


# ── EPUB 생성 ─────────────────────────────────────────────────────────────────

def _build_epub(req: ExportRequest) -> bytes:
    from ebooklib import epub

    book = epub.EpubBook()
    book.set_title(req.novel_title)
    book.set_language("ko")
    if req.author:
        book.add_author(req.author)

    # CSS
    css = epub.EpubItem(
        uid="style",
        file_name="style/main.css",
        media_type="text/css",
        content=textwrap.dedent("""\
            body { font-family: 'Malgun Gothic', 'NanumGothic', serif;
                   font-size: 1em; line-height: 1.9; margin: 1.5em 2em; }
            h1 { font-size: 1.6em; border-bottom: 1px solid #ccc;
                 padding-bottom: .3em; margin-bottom: 1em; }
            p  { margin: 0 0 .8em 0; text-indent: 1em; }
            blockquote { border-left: 3px solid #aaa; padding-left: 1em;
                         color: #555; margin: 1em 0; }
        """).encode("utf-8"),
    )
    book.add_item(css)

    chapters: list[epub.EpubHtml] = []
    spine = ["nav"]

    for ep in req.episodes:
        uid = f"ep{ep.episode_no}"
        chap = epub.EpubHtml(
            title=f"제{ep.episode_no}화 {ep.title}",
            file_name=f"{uid}.xhtml",
            lang="ko",
        )
        chap.content = (
            f'<html xmlns="http://www.w3.org/1999/xhtml">'
            f'<head><title>{ep.title}</title>'
            f'<link rel="stylesheet" href="style/main.css"/></head>'
            f'<body><h1>제{ep.episode_no}화&nbsp; {ep.title}</h1>'
            f'{ep.content_html or ""}'
            f'</body></html>'
        ).encode("utf-8")
        chap.add_item(css)
        book.add_item(chap)
        chapters.append(chap)
        spine.append(chap)

    book.toc = tuple(epub.Link(c.file_name, c.title, c.id) for c in chapters)
    book.spine = spine
    book.add_item(epub.EpubNcx())
    book.add_item(epub.EpubNav())

    buf = io.BytesIO()
    epub.write_epub(buf, book)
    return buf.getvalue()


# ── 메인 엔드포인트 ───────────────────────────────────────────────────────────

_MIME = {
    "txt":  ("text/plain; charset=utf-8",                          ".txt"),
    "md":   ("text/markdown; charset=utf-8",                       ".md"),
    "docx": ("application/vnd.openxmlformats-officedocument.wordprocessingml.document", ".docx"),
    "pdf":  ("application/pdf",                                    ".pdf"),
    "epub": ("application/epub+zip",                               ".epub"),
}

_BUILDERS = {
    "txt":  _build_txt,
    "md":   _build_md,
    "docx": _build_docx,
    "pdf":  _build_pdf,
    "epub": _build_epub,
}


def _safe_filename(name: str) -> str:
    """파일명에 쓸 수 없는 문자 제거"""
    return re.sub(r'[\\/:*?"<>|]', "_", name).strip() or "novel"


@router.post("/download")
async def export_download(req: ExportRequest):
    fmt = req.format.lower().strip()

    if fmt not in _BUILDERS:
        from fastapi import HTTPException
        raise HTTPException(status_code=400, detail=f"지원하지 않는 형식입니다: {fmt}")

    data = _BUILDERS[fmt](req)
    mime, ext = _MIME[fmt]
    filename = _safe_filename(req.novel_title) + ext

    # RFC 5987 인코딩 (한국어 파일명 처리)
    from urllib.parse import quote
    encoded_name = quote(filename, safe="")
    content_disposition = (
        f'attachment; filename="{filename}"; '
        f"filename*=UTF-8''{encoded_name}"
    )

    return StreamingResponse(
        io.BytesIO(data),
        media_type=mime,
        headers={"Content-Disposition": content_disposition},
    )
